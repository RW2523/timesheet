"""
DOCX parser — python-docx with Docling fallback for complex layouts.

Also exposes extract_embedded_images() for image-heavy / scanned DOCX files.
A DOCX is a ZIP archive; all embedded images live under word/media/.

Flow decision:
  1. Try Docling (handles tables/forms better).
  2. Fallback: python-docx paragraphs + tables.
  3. If very little text extracted → set ocr_required=True so the caller can
     pass the embedded images through OCR or VLM instead.
"""
import io
import logging
import os
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from app.core.config import settings

logger = logging.getLogger(__name__)

# Image extensions kept when extracting from the ZIP
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif", ".webp"}

# WMF/EMF are Windows metafiles — skip them, they break PIL
_SKIP_EXTS   = {".wmf", ".emf"}


class DocxParser:
    def parse(self, file_path: str) -> Dict[str, Any]:
        # ── 1. Try Docling first ───────────────────────────────────────────────
        try:
            from app.services.parsers.docling_parser import DoclingParser
            result = DoclingParser().parse(file_path)
            text   = result.get("raw_text") or ""
            if result and len(text.strip()) >= settings.MIN_TEXT_MEANINGFUL:
                return result
        except Exception as exc:
            logger.debug("Docling DOCX parse failed, falling back: %s", exc)

        # ── 2. python-docx ────────────────────────────────────────────────────
        try:
            from docx import Document
            doc = Document(file_path)

            paragraphs: List[str] = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            tables_data: List[Dict] = []
            for i, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    if any(c for c in row_cells):
                        rows.append(row_cells)
                if rows:
                    tables_data.append({"sheet": f"table_{i + 1}", "rows": rows})

            raw_text = "\n".join(paragraphs)
            if tables_data:
                raw_text += "\n\n" + "\n".join(
                    " | ".join(cell for cell in row)
                    for t in tables_data
                    for row in t["rows"]
                )

            total_chars = len(raw_text.strip())

            # ── 3. Low text → check for embedded images ────────────────────
            has_images, image_count = self._has_embedded_images(file_path)
            if total_chars < settings.MIN_TEXT_MEANINGFUL and not tables_data:
                return {
                    "raw_text":           raw_text or None,
                    "raw_tables":         None,
                    "ocr_required":       True,
                    "embedded_images":    image_count,
                    "confidence":         0.0,
                    "warnings": [
                        f"DOCX_LOW_TEXT: only {total_chars} chars — "
                        f"document contains {image_count} embedded image(s). "
                        "Use 'DOCX Image OCR' or 'DOCX VLM' parser to extract content."
                    ],
                    "extraction_method": "docx_low_text_ocr_pending",
                }

            return {
                "raw_text":          raw_text,
                "raw_tables":        tables_data,
                "ocr_required":      False,
                "embedded_images":   image_count,
                "confidence":        0.9,
                "warnings":          [],
                "extraction_method": "python-docx",
            }
        except Exception as exc:
            logger.error("DOCX parse error: %s", exc)
            return {
                "raw_text":          None,
                "raw_tables":        None,
                "ocr_required":      False,
                "embedded_images":   0,
                "confidence":        0.0,
                "warnings":          [str(exc)],
                "extraction_method": "docx_failed",
            }

    # ── Image helpers ──────────────────────────────────────────────────────────

    @staticmethod
    def _has_embedded_images(file_path: str) -> Tuple[bool, int]:
        """Return (has_images, count) without extracting anything."""
        count = 0
        try:
            with zipfile.ZipFile(file_path, "r") as zf:
                for name in zf.namelist():
                    if name.startswith("word/media/"):
                        ext = Path(name).suffix.lower()
                        if ext in _IMAGE_EXTS:
                            count += 1
        except Exception:
            pass
        return count > 0, count

    @staticmethod
    def extract_embedded_images(
        file_path: str,
        dest_dir: Optional[str] = None,
    ) -> List[str]:
        """Extract all raster images from a DOCX file to dest_dir.

        Returns a list of absolute file paths (one per image), sorted by name
        so page order is preserved.  If dest_dir is None a temp directory is
        created (caller is responsible for cleanup).
        """
        if dest_dir is None:
            dest_dir = tempfile.mkdtemp(prefix="docx_imgs_")

        paths: List[str] = []
        try:
            with zipfile.ZipFile(file_path, "r") as zf:
                media_names = sorted(
                    n for n in zf.namelist()
                    if n.startswith("word/media/")
                    and Path(n).suffix.lower() in _IMAGE_EXTS
                )
                for name in media_names:
                    data  = zf.read(name)
                    fname = os.path.basename(name)
                    out   = os.path.join(dest_dir, fname)
                    with open(out, "wb") as f:
                        f.write(data)
                    paths.append(out)
        except Exception as exc:
            logger.warning("DOCX image extraction failed: %s", exc)

        return paths
